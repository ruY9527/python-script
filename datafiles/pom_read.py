import os
import xml.etree.ElementTree as ET
import json
from docx import Document
from docx.shared import Cm

# --- 1. 定义命名空间 ---
# 确保命名空间字典正确
NS = {'mvn': 'http://maven.apache.org/POM/4.0.0'}

def load_metadata_map(file_path="metadata_map.json"):
    """加载自定义的组件元数据映射文件。"""
    mock_data = {
      "net.minidev:accessors-smart": {
        "country": "法国",
        "purpose": "反射增强(优化 JSON 序列化性能)",
        "is_donated": "否",
        "license_custom": "动态链接/Apache-2.0"
      },
      "com.fasterxml.jackson.core:jackson-databind": {
        "country": "美国",
        "purpose": "高性能 JSON 处理器",
        "is_donated": "否",
        "license_custom": "Apache-2.0 OR LGPL 2.1"
      }
    }
    try:
        # 使用 utf-8 编码打开文件，更安全
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"⚠️ 未找到元数据文件 '{file_path}'，将使用内置模拟数据。")
        return mock_data
    except Exception as e:
        print(f"❌ 加载元数据文件失败: {e}，将使用内置模拟数据。")
        return mock_data

# --- 2. 依赖提取函数 (重点修正) ---
def extract_dependencies(pom_file_path):
    dependencies = set()
    properties = {}
    
    try:
        # 强制使用 utf-8 编码读取文件，解决编码问题
        with open(pom_file_path, 'r', encoding='utf-8') as f:
            pom_content = f.read()
            root = ET.fromstring(pom_content)

        # 1. 提取属性 (Properties)
        properties_element = root.find('mvn:properties', NS)
        if properties_element is not None:
            for prop in properties_element:
                # prop.tag 应该形如 {http://maven.apache.org/POM/4.0.0}propertyName
                tag_name = ET.QName(prop.tag).localname # 提取本地标签名 (e.g., spring.version)
                properties[tag_name] = prop.text.strip() if prop.text else ''

        # 2. 查找 <dependencies> 块并解析
        dependencies_element = root.find('mvn:dependencies', NS)
        
        if dependencies_element is not None:
            for dep in dependencies_element.findall('mvn:dependency', NS):
                # 使用 findtext 简化非空检查和文本提取
                g = dep.findtext('mvn:groupId', default='', namespaces=NS).strip()
                a = dep.findtext('mvn:artifactId', default='', namespaces=NS).strip()
                v_raw = dep.findtext('mvn:version', default='', namespaces=NS).strip()

                if g and a:
                    v = v_raw
                    # 3. 替换属性占位符
                    if v_raw.startswith('${') and v_raw.endswith('}'):
                        prop_name = v_raw[2:-1] 
                        # 查找属性，找不到则保留占位符
                        v = properties.get(prop_name, v_raw) 

                    dependencies.add((g, a, v))
                    
    except ET.ParseError as e:
        print(f"⚠️ 无法解析 XML 文件 {pom_file_path} (XML结构错误或编码问题): {e}")
    except FileNotFoundError:
        print(f"❌ 文件未找到: {pom_file_path}")
    except Exception as e:
        print(f"❌ 处理文件 {pom_file_path} 时发生错误: {e}")
        
    return dependencies

# --- 3. 报告生成函数 (与上一个版本基本一致，确保逻辑正确) ---
def generate_dependency_report(root_dir, output_file="dependency_audit_report.docx", metadata_map={}):
    all_unique_dependencies = set()
    found_pom_count = 0
    
    print(f"🚀 开始遍历目录：{root_dir}")

    # 使用 os.path.exists 检查根目录是否存在
    if not os.path.exists(root_dir):
        print(f"🚨 错误：指定的项目路径 '{root_dir}' 不存在。请检查路径是否正确。")
        return

    for dirpath, _, filenames in os.walk(root_dir):
        if 'pom.xml' in filenames:
            pom_path = os.path.join(dirpath, 'pom.xml')
            print(f"    - 发现并解析：{pom_path}")
            found_pom_count += 1
            all_unique_dependencies.update(extract_dependencies(pom_path))
            
    # 检查点 1: 是否找到了 pom 文件
    if found_pom_count == 0:
        print(f"\n🚨 警告：在路径 '{root_dir}' 中未找到任何 'pom.xml' 文件。")
        return
        
    # 检查点 2: 是否提取到了依赖
    if not all_unique_dependencies:
        print(f"🚨 警告：找到了 {found_pom_count} 个 pom 文件，但未提取到任何实际的依赖。Word 文档表格将为空。")
        # 即使为空也继续生成文档，以便用户看到文档结构

    print(f"\n✅ 依赖解析完成。共发现 {len(all_unique_dependencies)} 个不重复依赖。")
    
    # --- 生成 Word 文档 ---
    document = Document()
    document.add_heading('Java 项目依赖分析报告', 0)
    
    headers = [
        '序号', '类型 (组件/软件/项目/其他)', '作者/项目/版本', '国别', 
        '仓库地址', '用途', '是否捐赠基金会', '对应事项 (引入方式/调用方式和许可证等)'
    ]
    
    # rows = 依赖数量 + 1 (表头)
    table = document.add_table(rows=len(all_unique_dependencies) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 填充表头
    for i, header_text in enumerate(headers):
        table.rows[0].cells[i].text = header_text
        
    # 填充数据行
    sorted_dependencies = sorted(list(all_unique_dependencies))
    for i, dep in enumerate(sorted_dependencies, 1):
        group_id, artifact_id, version = dep
        
        map_key = f"{group_id}:{artifact_id}"
        # 使用 get 确保找不到时返回通用占位符
        extra_info = metadata_map.get(map_key, {
            'country': '待补充',
            'purpose': '待补充',
            'is_donated': '否',
            'license_custom': '待补充/开源组件许可证'
        })
        
        row_cells = table.rows[i].cells
        
        row_cells[0].text = str(i)
        row_cells[1].text = '开源组件'
        
        # 兼容您的表格格式：artifactId/version (并在下方添加 GAV 坐标)
        project_version_text = f'{artifact_id}/{version}'
        gav_detail = f'({group_id}:{artifact_id}:{version})'
        
        # ⚠️ 如果版本未解析，进行标记
        if version.startswith('${') and version.endswith('}'):
             project_version_text += " (版本未解析)"
             gav_detail = "(父POM或外部属性未解析，需手动检查)"

        row_cells[2].text = f'{project_version_text}\n{gav_detail}'
        
        # 从映射表填充数据
        row_cells[3].text = extra_info['country']
        row_cells[4].text = f'Maven Central (或 {group_id})'
        row_cells[5].text = extra_info['purpose']
        row_cells[6].text = extra_info['is_donated']
        row_cells[7].text = extra_info['license_custom']

    document.save(output_file)
    print(f"\n🎉 报告已成功生成到文件: {os.path.abspath(output_file)}")

# --- 4. 运行代码示例 ---

if __name__ == '__main__':
    # ⚠️ 再次确认此路径是您本地文件系统的有效绝对或相对路径
    project_path = "/Users/baoyang/Documents/coding/yundee_coding/dbs_platform"
    
    # 1. 加载元数据映射表
    metadata = load_metadata_map()
    
    # 2. 生成报告
    generate_dependency_report(project_path, "dependency_audit_report.docx", metadata)